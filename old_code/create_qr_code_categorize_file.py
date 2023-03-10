import pandas as pd
import numpy as np
import math
import os
import re
import pyqrcode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import pytesseract as pt
from PyPDF2 import PdfFileReader, PdfFileWriter
# from PIL import Image
from pdf2image import convert_from_path
import shutil
from docx2pdf import convert

pt.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'

root = 'D:/Shweta/data_digitization'
master_file = pd.read_excel(os.path.join(root, 'reference_docs/2022_03_19_patient_master_list_sk.xlsx'))
report_names_df = pd.read_excel(os.path.join(root, 'reference_docs/Report_types_17.xlsx'))
categorized_files_df = pd.read_excel(os.path.join(root, 'reference_docs/2010_file_categorization_excel.xlsx'))
scanned_patient_file_path = os.path.join(root, 'scanned_patient_files/2022_03_14')
categorized_file_path = os.path.join(root, 'scanned_patient_files/2022_03_14/original_pdf')

def change_sep(string, old_sep, new_sep):
    """ change the separator between the string or within the string
    :param string: string
    :param old_sep: string separators (' ', '_', '/')
    :param new_sep: string separators (' ', '_', '/')
    :return: string with changed separators
    """
    changed_sep = re.sub(old_sep, new_sep, str(string))
    return changed_sep

id_cols = ['file_number', 'MR number', 'DATE']

def get_id_data(master_list, index, id_cols):
    """
    get id values from input id names and single row of master list
    :param master_list: pd.DataFrame
    :param index: int
    :return: list of id data for single row of master list
    """
    id_dat = []
    for id_col in id_cols:
        id = master_list[id_col][index]
        id_dat.append(id)
    return id_dat

folder_col_heads = ['Report Name', 'Subfolder Name']

def get_folder_subfolder(categorized_excel, index):
    """
    it will give the folder name and sub-folder name for the qr code
    :param categorized_excel: pd.DataFrame categorized excel
    :param index: integer
    :return:
    """
    folder_dat = []
    for col_name in folder_col_heads:
        folder_info = categorized_excel[col_name][index]
        folder_dat.append(folder_info)
    return folder_dat

report_types = ['Patient Information', 'Clinical Examination', 'Radiology', 'Metastatic Examination', 'Biopsy Pathology',
                'Neo-Adjuvant Chemotherapy', 'Surgical Procedures', 'Patient Images', 'Surgery Media', 'Surgery Pathology',
                'Chemotherapy', 'Radiotherapy', 'Follow-up Notes', 'Genetics', 'Miscellaneous', 'Patient File Data', 'PROMS']


def make_qr_code(file_number, mr_number, report_type, subfolder, destination):
    file_number_str = re.sub('_', '/', str(file_number))
    if subfolder is not None:
        qr_code = file_number_str + '_' + \
            str(mr_number) + '_' + str(report_type) + '_' + str(subfolder)
    else:
        qr_code = file_number_str + '_' + \
            str(mr_number) + '_' + str(report_type)
    qr = pyqrcode.create(qr_code)
    report_type_for_name = re.sub(' ', '_', str(report_type))
    qr_img_name = file_number + '_' + \
        str(mr_number) + '_' + report_type_for_name + '.png'
    qr_path = os.path.join(destination, qr_img_name)
    qr.png(qr_path, scale=4)
    print('QR code created for ' + file_number + ' ' + report_type + ' ')
    return qr_img_name

def format_word_doc(doc, id_value):
    text = doc.add_paragraph()
    report_type_name = text.add_run(str(id_value))
    report_type_name.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    report_type_name.bold = True
    report_type_name.font.size = Pt(28)
    report_type_name.font.name = 'Arial Black'


def add_qr_code_in_word_doc(report_type, qr_code_path, file_number, mr_number, patient_name, dob, tmp_folder_path):
    doc = Document()
    doc.add_picture(qr_code_path)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    type = report_type
    text = doc.add_paragraph()
    report_type_name = text.add_run(str(type))
    report_type_name.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    report_type_name.bold = True
    report_type_name.font.size = Pt(28)
    report_type_name.font.name = 'Arial Black'
    blank_para = doc.add_paragraph()
    run = blank_para.add_run()
    run.add_break()
    file_number_str = re.sub('_', '/', str(file_number))
    file_no = 'File Number: ' + str(file_number_str)
    id = doc.add_paragraph().add_run(file_no)
    id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    id.font.size = Pt(20)
    id.font.name = 'Arial Black'
    mr_no = 'MR Number: ' + str(mr_number)
    id = doc.add_paragraph().add_run(mr_no)
    id.font.size = Pt(20)
    id.font.name = 'Arial Black'
    pt_name = 'Patient Name: ' + patient_name
    id = doc.add_paragraph().add_run(pt_name)
    id.font.size = Pt(20)
    id.font.name = 'Arial Black'
    date_of_birth = 'Date of Birth: ' + dob
    id = doc.add_paragraph().add_run(date_of_birth)
    id.font.size = Pt(20)
    id.font.name = 'Arial Black'
    report_type = re.sub(' ', '_', report_type)
    report_type = report_type.lower()
    doc_name = str(file_number) + '_' + str(mr_number) + \
        '_' + str(report_type) + '.docx'
    coded_data = os.path.join(tmp_folder_path, 'coded_data')
    if not os.path.isdir(coded_data):
        os.mkdir(coded_data)
    doc_path = os.path.join(coded_data, doc_name)
    doc.save(doc_path)
    return doc_path

def convert_doc_to_pdf(doc_path):
    pdf_path = re.sub('.docx', '.pdf', str(doc_path))
    convert(doc_path, pdf_path)
    return pdf_path


# pdf_path = convert_doc_to_pdf(
#     'D:/Shweta/data_digitization/sample_output/2022_03_15/coded_data/38_10_302_01_patient_information.docx')

# add_qr_code_in_word_doc(report_type = 'patient_information',
#                 qr_code_path = 'D:/Shweta/data_digitization/sample_output/2022_03_14/qr_codes/38_10_302_01_patient_information.png'
#                         , file_number = '12_13', mr_number = '1213', patient_name='xyz', dob='dmy',
#                         tmp_folder_path = 'D:/Shweta/data_digitization/sample_output/2022_03_14')

def split_pdf_to_pages(file_number, scanned_files_path, splitted_file_path):
    pdf_file_name = str(file_number) + '.pdf'
    scanned_file_path = os.path.join(scanned_files_path, pdf_file_name)
    pdf_file = PdfFileReader(scanned_file_path)
    pages = pdf_file.get_pages
    pages = convert_from_path(scanned_file_path, 500,
                              poppler_path='C:/Program Files/poppler-0.68.0/bin')
    i = 0
    for index, page in enumerate(pages):
        if i == index:
            page_no = i + 1
            out_jpg = str(file_number) + '_' + str(page_no) + '.jpg'
            page.save(os.path.join(splitted_file_path, out_jpg), 'JPEG')
            i += 1
    print("file number: ", file_number + " split")

def split_pdf_by_pages(file_number, scanned_files_path, splitted_file_path):
    pdf_file_name = str(file_number) + '.pdf'
    scanned_file = PdfFileReader(os.path.join(scanned_files_path, pdf_file_name))
    page_range = scanned_file.getNumPages()
    for i in range(page_range):
        page = scanned_file.getPage(i)
        page_no = i + 1
        splitted_file = str(file_number) + '_' + str(page_no) + '.pdf'
        pdf_writer = PdfFileWriter()
        pdf_writer.addPage(page)
        with open(os.path.join(splitted_file_path, splitted_file), 'wb') as out:
            pdf_writer.write(out)
    print("file number: ", file_number + " splitted")

def get_image_no(file_number, file_images_lst):
    file_images_no_lst = []
    for file_image in file_images_lst:
        file_image_no = re.sub(file_number, '', str(file_image))
        file_image_no = re.sub('.jpg', '', file_image_no)
        file_image_no = re.sub('_', '', file_image_no)
        file_image_no = file_image_no.strip()
        file_images_no_lst.append(file_image_no)
    return file_images_no_lst

def split_report_page_no(report_page_no):
    if isinstance(report_page_no, float):
        page_no_lst = []
        if not math.isnan(report_page_no):
            integer = int(report_page_no)
            page_no_lst.append(str(integer))
        return page_no_lst
    elif isinstance(report_page_no, int):
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    elif ';' in report_page_no:
        page_no_lst = []
        report_page_no_splitted = report_page_no.split(';')
        for page_no in report_page_no_splitted:
            if '|' in page_no:
                partitions = page_no.partition('|')
                start = int(partitions[0])
                end = int(partitions[2]) + 1
                page_nos = np.arange(start, end)
                page_nos_lst = page_nos.tolist()
                for no in page_nos_lst:
                    page_no_lst.append(str(no))
            else:
                page_no_lst.append(str(page_no))
        return page_no_lst
    elif '|' in report_page_no:
        page_no_lst = []
        partitions = report_page_no.partition('|')
        start = int(partitions[0])
        end = int(partitions[2]) + 1
        page_nos = np.arange(start, end)
        page_nos_lst = page_nos.tolist()
        for no in page_nos_lst:
            page_no_lst.append(str(no))
        return page_no_lst
    elif type(report_page_no) in (float, int):
        page_no_lst = []
        report_page_no = int(report_page_no)
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    else:
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst


def classify_file_images_by_report_types(splitted_scanned_file_path, report_page_nums, file_number, report_type, destination_path):
    splitted_scanned_files = os.listdir(splitted_scanned_file_path)
    img_no_lst = get_image_no(file_number, splitted_scanned_files)
    report_page_no_splitted = split_report_page_no(report_page_nums)
    file_no_dir = os.path.join(destination_path, str(file_number))
    if not os.path.isdir(file_no_dir):
        os.mkdir(file_no_dir)
    report_dir = os.path.join(file_no_dir, report_type)
    if not os.path.isdir(report_dir):
        os.mkdir(report_dir)
    for page_no in report_page_no_splitted:
        if page_no in img_no_lst:
            report_page_name = str(file_number) + '_' + str(page_no) + '.jpg'
            source_path = os.path.join(
                splitted_scanned_file_path, report_page_name)
            dest_path = os.path.join(report_dir, report_page_name)
            shutil.copy(source_path, dest_path)


classify_file_images_by_report_types('D:/Shweta/data_digitization/scanned_patient_files/2022_03_15/splitted_pdf/38_10',
                                     '2;47;48;50|52;54;55', '38_10', '01_patient_information',
                                     'D:/Shweta/data_digitization/scanned_patient_files/2022_03_15/categorized_imgs')


def rename_images(pdf_doc_path, dir_path, file_no, report_type, destination_path):
    report_dir = os.path.join(dir_path, str(report_type))
    img_list = os.listdir(report_dir)
    for index, img in enumerate(img_list):
        old_file_path = os.path.join(report_dir, img)
        img_no = index + 1
        new_name = str(file_no) + '_' + str(img_no) + '.jpg'
        file_dir = os.path.join(destination_path, str(file_no))
        if not os.path.isdir(file_dir):
            os.mkdir(file_dir)
        new_file_path = os.path.join(file_dir, report_type)
        if not os.path.isdir(new_file_path):
            os.mkdir(new_file_path)
        dest_path = os.path.join(new_file_path, new_name)
        shutil.copy(old_file_path, dest_path)
        coded_file_name = 'code_' + \
            str(file_no) + '_' + str(report_type) + '.pdf'
        shutil.copy(pdf_doc_path, os.path.join(new_file_path, coded_file_name))
        print('report_renamed')


rename_images(pdf_path, 'D:/Shweta/data_digitization/scanned_patient_files/2022_03_15/categorized_imgs/38_10',
              '38_10', '01_patient_information',
              'D:/Shweta/data_digitization/scanned_patient_files/2022_03_15/classified_and_renamed')

# classify_file_images_by_report_types('D:/Shweta/data_digitization/sample_output/2022_03_14/splitted_files/38_10',
#                                      '2;47;48;50|52;54;55', '38_10', '01_patient_information',
#                     'D:/Shweta/data_digitization/scanned_patient_files/2022_03_14/categorrized_and_renamed_files')


def categorize_file_by_report_types(report_names_df, categorized_files_df, splitted_files_path, tmp_folder_path, destination_path):
    for i in range(len(categorized_files_df)):
        file_number = categorized_files_df['file_number'][i]
        mr_number = categorized_files_df['mr_number'][i]
        patient_name = categorized_files_df['patient_name'][i]
        dob = categorized_files_df['date_of_birth'][i]
        print(file_number, mr_number, patient_name, dob)
        for report_type in report_names_df['report_number_and_type']:
            print(report_type)
            qr_code_dir = os.path.join(tmp_folder_path, 'qr_codes')
            if not os.path.isdir(qr_code_dir):
                os.mkdir(qr_code_dir)
            qr_img_name = make_qr_code(
                file_number, mr_number, report_type, None, qr_code_dir)
            coded_data_dir = os.path.join(tmp_folder_path, 'coded_data')
            if not os.path.isdir(coded_data_dir):
                os.mkdir(coded_data_dir)
            qr_code_path = os.path.join(qr_code_dir, qr_img_name)
            coded_doc_path = add_qr_code_in_word_doc(report_type, qr_code_path, file_number, mr_number, patient_name, dob,
                                                     tmp_folder_path)
            coded_pdf_path = convert_doc_to_pdf(coded_doc_path)
            report_type_str = re.sub(' ', '_', str(report_type))
            report_page_nums = categorized_files_df[report_type_str][i]
            splitted_images_for_file_no = os.path.join(
                splitted_files_path, str(file_number))
            classified_files_path = os.path.join(
                tmp_folder_path, 'classfied_files')
            print(classified_files_path)
            if not os.path.isdir(classified_files_path):
                os.mkdir(classified_files_path)
            classify_file_images_by_report_types(splitted_images_for_file_no, str(report_page_nums), file_number, report_type_str,
                                                 classified_files_path)
            renamed_files_path = os.path.join(
                classified_files_path, str(file_number))
            print(renamed_files_path)
            rename_images(renamed_files_path, str(file_number),
                          report_type_str, destination_path)
            print('file: ' + file_number +
                  ' classified by report types and arranged by sequence')


categorize_file_by_report_types(report_names_df, categorized_files_df,
                                'D:/Shweta/data_digitization/scanned_patient_files/2022_03_15/splitted_pdf',
                                'D:/Shweta/data_digitization/sample_output/2022_03_15',
                                'D:/Shweta/data_digitization/sample_output/2022_03_15/destination_path')


def split_report_page_no(report_page_no):
    if isinstance(report_page_no, float):
        page_no_lst = []
        if not math.isnan(report_page_no):
            integer = int(report_page_no)
            page_no_lst.append(str(integer))
        return page_no_lst
    elif isinstance(report_page_no, int):
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    elif ';' in report_page_no:
        page_no_lst = []
        report_page_no_splitted = report_page_no.split(';')
        for page_no in report_page_no_splitted:
            if '|' in page_no:
                partitions = page_no.partition('|')
                start = int(partitions[0])
                end = int(partitions[2]) + 1
                if start<end:
                    page_nos = np.arange(start, end)
                    page_nos_lst = page_nos.tolist()
                    for no in page_nos_lst:
                        page_no_lst.append(str(no))
                else:
                    page_nos = np.arange(end, start)
                    page_nos = page_nos[::-1]
                    page_nos_lst = page_nos.tolist()
                    for no in page_nos_lst:
                        page_no_lst.append(str(no))
            else:
                page_no_lst.append(str(page_no))
        return page_no_lst
    elif '|' in report_page_no:
        page_no_lst = []
        partitions = report_page_no.partition('|')
        start = int(partitions[0])
        end = int(partitions[2]) + 1
        page_nos = np.arange(start, end)
        page_nos_lst = page_nos.tolist()
        for no in page_nos_lst:
            page_no_lst.append(str(no))
        return page_no_lst
    elif type(report_page_no) in (float, int):
        page_no_lst = []
        report_page_no = int(report_page_no)
        page_no_lst.append(str(report_page_no))
        return page_no_lst
    else:
        page_no_lst = []
        page_no_lst.append(str(report_page_no))
        return page_no_lst

